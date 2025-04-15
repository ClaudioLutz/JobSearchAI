import os
import json
import logging
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("word_template_generator.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("word_template_generator")

def json_to_docx(motivation_letter_json, output_path=None):
    """
    Convert motivation letter JSON to a Word document
    
    Args:
        motivation_letter_json (dict): The motivation letter data in JSON format
        output_path (str, optional): The path to save the Word document. If None, a default path will be generated.
        
    Returns:
        str: The path to the saved Word document
    """
    try:
        logger.info("Creating Word document from JSON data")
        
        # Create a new Document
        doc = Document()
        
        # Set up the document margins (2.5 cm on all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Extract fields from JSON
        candidate_name = motivation_letter_json.get('candidate_name', '')
        candidate_address = motivation_letter_json.get('candidate_address', '')
        candidate_city = motivation_letter_json.get('candidate_city', '')
        candidate_email = motivation_letter_json.get('candidate_email', '')
        candidate_phone = motivation_letter_json.get('candidate_phone', '')
        
        company_name = motivation_letter_json.get('company_name', '')
        company_department = motivation_letter_json.get('company_department', '')
        company_address = motivation_letter_json.get('company_address', '')
        
        date = motivation_letter_json.get('date', '')
        subject = motivation_letter_json.get('subject', '')
        greeting = motivation_letter_json.get('greeting', '')
        introduction = motivation_letter_json.get('introduction', '')
        
        body_paragraphs = motivation_letter_json.get('body_paragraphs', [])
        
        closing = motivation_letter_json.get('closing', '')
        signature = motivation_letter_json.get('signature', '')
        full_name = motivation_letter_json.get('full_name', '')
        
        # Add candidate information
        p = doc.add_paragraph()
        p.add_run(candidate_name).bold = True
        p.add_run('\n' + candidate_address)
        p.add_run('\n' + candidate_city)
        p.add_run('\n' + candidate_email)
        p.add_run('\n' + candidate_phone)
        
        # Add some space
        doc.add_paragraph()
        
        # Add company information
        p = doc.add_paragraph()
        p.add_run(company_name)
        if company_department:
            p.add_run('\n' + company_department)
        if company_address:
            p.add_run('\n' + company_address)
        
        # Add date
        p = doc.add_paragraph()
        p.add_run(date)
        
        # Add some space
        doc.add_paragraph()
        
        # Add subject line
        p = doc.add_paragraph()
        p.add_run(subject).bold = True
        
        # Add some space
        doc.add_paragraph()
        
        # Add greeting
        p = doc.add_paragraph()
        p.add_run(greeting + ',')
        
        # Add introduction
        p = doc.add_paragraph()
        p.add_run(introduction)
        
        # Add body paragraphs
        for paragraph in body_paragraphs:
            p = doc.add_paragraph()
            p.add_run(paragraph)
        
        # Add closing
        p = doc.add_paragraph()
        p.add_run(closing)
        
        # Add signature
        p = doc.add_paragraph()
        p.add_run(signature + ',')
        p.add_run('\n' + full_name)
        
        # If no output path is provided, generate one
        if not output_path:
            # Create the motivation letters directory if it doesn't exist
            motivation_letters_dir = Path('motivation_letters')
            motivation_letters_dir.mkdir(exist_ok=True)
            
            # Generate a filename based on the job title
            # Extract job title from subject line
            job_title = subject.replace('Bewerbung als ', '').replace('Bewerbung für ', '')
            job_title = ''.join(c if c.isalnum() or c in [' ', '_', '-'] else '_' for c in job_title)
            job_title = job_title.replace(' ', '_')[:30]
            
            # Create the filename - use job title as the primary identifier
            filename = f"motivation_letter_{job_title}.docx"
            output_path = str(motivation_letters_dir / filename)
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Word document saved to: {output_path}")
        
        return output_path
    
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return None

def create_word_document_from_json_file(json_file_path):
    """
    Create a Word document from a JSON file
    
    Args:
        json_file_path (str): The path to the JSON file
        
    Returns:
        str: The path to the saved Word document
    """
    try:
        logger.info(f"Creating Word document from JSON file: {json_file_path}")
        
        # Load the JSON data
        with open(json_file_path, 'r', encoding='utf-8') as f:
            motivation_letter_json = json.load(f)
        
        # Generate the output path
        output_path = json_file_path.replace('.json', '.docx')
        
        # Create the Word document
        return json_to_docx(motivation_letter_json, output_path)
    
    except Exception as e:
        logger.error(f"Error creating Word document from JSON file: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return None

if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) > 1:
        json_file_path = sys.argv[1]
        docx_path = create_word_document_from_json_file(json_file_path)
        
        if docx_path:
            print(f"Word document created successfully: {docx_path}")
        else:
            print("Failed to create Word document")
    else:
        print("Usage: python word_template_generator.py <json_file_path>")
